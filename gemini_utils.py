from __future__ import annotations

import json
import os
import io
import google.genai as genai
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple
import streamlit as st
from PIL import Image
import cv2
import numpy as np
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Load API key from environment
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyCcPgyoaVlx0FDjsbPbPJ0OTo1xlahjcUk")
client = genai.Client(api_key=GEMINI_API_KEY)

DEFAULT_MODEL = "gemini-2.5-flash"
SUPPORTED_FILE_TYPES = ["txt", "docx", "pdf", "jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff"]
SUPPORTED_TARGET_LANGUAGES = ["English", "Arabic", "French", "Spanish", "German", "Chinese", "Japanese", "Korean", "Russian", "Portuguese", "Italian", "Dutch", "Turkish", "Persian", "Urdu", "Hindi"]

HISTORY_DIR = Path("history_store")
HISTORY_DIR.mkdir(exist_ok=True)


def _safe_name(name: str) -> str:
    return "".join(ch if ch.isalnum() or ch in "-_." else "_" for ch in name)


def save_history_record(app_name: str, file_name: str, payload: Dict[str, Any]) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = HISTORY_DIR / f"{app_name}_{ts}_{_safe_name(file_name)}.json"
    record = {
        "app_name": app_name,
        "file_name": file_name,
        "saved_at": datetime.now().isoformat(),
        "payload": payload,
    }
    path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def list_history_records(prefix: str | None = None) -> List[Path]:
    items = sorted(HISTORY_DIR.glob("*.json"), reverse=True)
    if prefix:
        items = [p for p in items if p.name.startswith(prefix)]
    return items


def load_history_record(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def clear_history(prefix: str | None = None) -> None:
    for item in list_history_records(prefix=prefix):
        try:
            item.unlink()
        except OSError:
            pass


def file_to_bytes(uploaded_file) -> bytes:
    """Convert uploaded file to bytes."""
    return uploaded_file.read()


def build_translation_context(file_name: str, extracted_text: str, target_language: str) -> str:
    """Build context for translation."""
    return f"""
    File: {file_name}
    Target Language: {target_language}
    
    Extracted Text:
    {extracted_text}
    
    Please provide a professional translation of the above text to {target_language}.
    Maintain the original formatting and structure as much as possible.
    """


def call_gemini_for_translation(filename: str, file_bytes: bytes, target_language: str, model_name: str = DEFAULT_MODEL, deep_scan: bool = False) -> Dict[str, Any]:
    """Call Gemini API for translation with file processing."""
    try:
        # Extract text from file
        extracted_text = extract_text_from_file(file_bytes, filename, model_name)
        
        if "error" in extracted_text.lower():
            return {
                "filename": filename,
                "target_language": target_language,
                "extracted_text": extracted_text,
                "translation": f"Extraction error: {extracted_text}",
                "model_used": model_name,
                "deep_scan": deep_scan,
                "success": False,
                "detected_language": "Unknown",
                "original_text": extracted_text,
                "translated_text": f"Extraction error: {extracted_text}",
                "examples": []
            }
        
        # Detect language (simple heuristic)
        detected_language = "Unknown"
        if any(char in extracted_text for char in ['ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ']):
            detected_language = "Arabic"
        elif any(char in extracted_text for char in ['α', 'β', 'γ', 'δ', 'ε', 'ζ', 'η', 'θ']):
            detected_language = "Greek"
        elif any(char in extracted_text for char in ['あ', 'い', 'う', 'え', 'お', 'か', 'き', 'く']):
            detected_language = "Japanese"
        elif any(char in extracted_text for char in ['한', '글', '이', '가', '을', '는']):
            detected_language = "Korean"
        elif any(char in extracted_text for char in ['中', '文', '字', '语']):
            detected_language = "Chinese"
        elif any(char in extracted_text for char in ['ß', 'ä', 'ö', 'ü', 'Ä', 'Ö', 'Ü']):
            detected_language = "German"
        elif any(char in extracted_text for char in ['à', 'é', 'è', 'ê', 'ë', 'ç', 'ù', 'û']):
            detected_language = "French"
        elif any(char in extracted_text for char in ['ñ', 'ñ', 'ü', 'ü', '¿', '¡']):
            detected_language = "Spanish"
        else:
            detected_language = "English"
        
        # Translate the extracted text
        prompt = f"""
        Translate the following text to {target_language}:
        
        {extracted_text}
        
        Provide only the translation without additional commentary. Maintain the original formatting and structure as much as possible.
        """
        
        response = client.models.generate_content(
            model=model_name,
            contents=prompt
        )
        translation = response.text
        
        # Generate examples
        examples_prompt = f"""
        Based on this translation from {detected_language} to {target_language}, provide 2-3 simple examples in layman English that demonstrate the key concepts or phrases:
        
        Original: {extracted_text[:500]}...
        Translation: {translation[:500]}...
        
        Provide only 2-3 brief examples, numbered.
        """
        
        try:
            examples_response = client.models.generate_content(
                model=model_name,
                contents=examples_prompt
            )
            examples_text = examples_response.text
            examples = [line.strip() for line in examples_text.split('\n') if line.strip() and any(char.isdigit() for char in line)][:3]
        except:
            examples = []
        
        return {
            "filename": filename,
            "target_language": target_language,
            "extracted_text": extracted_text,
            "translation": translation,
            "model_used": model_name,
            "deep_scan": deep_scan,
            "success": True,
            "detected_language": detected_language,
            "original_text": extracted_text,
            "translated_text": translation,
            "examples": examples
        }
        
    except Exception as e:
        return {
            "filename": filename,
            "target_language": target_language,
            "extracted_text": "",
            "translation": f"Translation error: {str(e)}",
            "model_used": model_name,
            "deep_scan": deep_scan,
            "success": False,
            "detected_language": "Unknown",
            "original_text": "",
            "translated_text": f"Translation error: {str(e)}",
            "examples": []
        }


def chat_with_context(question: str, context: str, model: str = DEFAULT_MODEL) -> str:
    """Chat with Gemini using provided context."""
    try:
        prompt = f"""
        Context: {context}
        
        Question: {question}
        
        Please answer the question based on the provided context. Be helpful and specific.
        """
        
        response = client.models.generate_content(
            model=model,
            contents=prompt
        )
        return response.text
    except Exception as e:
        return f"Chat error: {str(e)}"


def build_docx_bytes(file_name: str, translated_content: str) -> bytes:
    """Build DOCX file from translated content."""
    doc = Document()
    doc.add_heading(f"Translation of {file_name}", 0)
    doc.add_paragraph(translated_content)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_json_bytes(translated_content: str) -> bytes:
    """Build JSON file from translated content."""
    data = {
        "translation": translated_content,
        "timestamp": datetime.now().isoformat()
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def render_diagram_explanations(image_bytes: bytes, model: str = DEFAULT_MODEL) -> str:
    """Generate explanations for diagrams/images using Gemini Vision."""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        
        prompt = """
        Analyze this image and provide:
        1. A detailed description of what you see
        2. Any text content visible in the image
        3. The type of diagram or visual content (chart, table, flowchart, etc.)
        4. Key information that should be preserved in translation
        
        Be thorough and specific.
        """
        
        response = client.models.generate_content(
            model=model,
            contents=[prompt, image]
        )
        return response.text
    except Exception as e:
        return f"Image analysis error: {str(e)}"


def extract_text_from_file(file_bytes: bytes, file_name: str, model: str = DEFAULT_MODEL) -> str:
    """Extract text from various file types using Gemini."""
    try:
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension in [".txt"]:
            return file_bytes.decode("utf-8", errors="replace")
        
        elif file_extension in [".docx"]:
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file_extension in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"]:
            image = Image.open(io.BytesIO(file_bytes))
            
            prompt = "Extract all text from this image. Preserve the structure and formatting as much as possible."
            response = client.models.generate_content(
                model=model,
                contents=[prompt, image]
            )
            return response.text
        
        elif file_extension == ".pdf":
            # For PDF, we'll need PyMuPDF
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                return text
            except ImportError:
                return "PDF processing requires PyMuPDF. Please install it with: pip install PyMuPDF"
        
        else:
            return f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        return f"Text extraction error: {str(e)}"
